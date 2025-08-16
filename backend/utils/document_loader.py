"""
Document Loader for Industry Reporter 2
Enhanced document loading with support for multiple file formats
"""
import os
import json
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional, Set
from datetime import datetime

# Document processing imports
import PyPDF2
import docx
import pandas as pd
from bs4 import BeautifulSoup
import chardet

from core.config import config
from core.logging import get_logger

logger = get_logger(__name__)


class DocumentLoader:
    """
    Enhanced document loader with support for multiple file formats
    """
    
    def __init__(self, base_path: str = None, **kwargs):
        self.base_path = base_path or config.settings.doc_path
        
        # Configuration
        self.max_file_size_mb = kwargs.get('max_file_size_mb', config.settings.max_doc_size_mb)
        self.supported_formats = kwargs.get('supported_formats', config.settings.supported_formats)
        self.recursive = kwargs.get('recursive', True)
        self.include_hidden = kwargs.get('include_hidden', False)
        
        # Processing options
        self.extract_metadata = kwargs.get('extract_metadata', True)
        self.preserve_structure = kwargs.get('preserve_structure', True)
        self.chunk_large_files = kwargs.get('chunk_large_files', True)
        self.max_chunk_size = kwargs.get('max_chunk_size', 50000)  # characters
        
        # Content filtering
        self.min_content_length = kwargs.get('min_content_length', 100)
        self.exclude_patterns = kwargs.get('exclude_patterns', ['.git', '__pycache__', '.env'])
        
        # Encoding detection
        self.default_encoding = kwargs.get('default_encoding', 'utf-8')
        self.detect_encoding = kwargs.get('detect_encoding', True)
        
        # Statistics
        self._stats = {
            'files_processed': 0,
            'files_skipped': 0,
            'total_size_bytes': 0,
            'processing_time': 0,
            'errors': []
        }
    
    async def load(self) -> List[Dict[str, Any]]:
        """Load all supported documents from the base path"""
        try:
            start_time = datetime.now()
            self._reset_stats()
            
            if not os.path.exists(self.base_path):
                logger.warning(f"Document path does not exist: {self.base_path}")
                return []
            
            # Find all supported files
            file_paths = await self._find_supported_files()
            
            if not file_paths:
                logger.info(f"No supported documents found in: {self.base_path}")
                return []
            
            # Process files concurrently
            documents = await self._process_files_batch(file_paths)
            
            # Update statistics
            end_time = datetime.now()
            self._stats['processing_time'] = (end_time - start_time).total_seconds()
            
            logger.info(f"Loaded {len(documents)} documents in {self._stats['processing_time']:.2f}s")
            
            return documents
            
        except Exception as e:
            logger.error(f"Document loading failed: {e}")
            self._stats['errors'].append(str(e))
            return []
    
    async def load_file(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Load a single document file"""
        try:
            if not os.path.exists(file_path):
                return None
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size_mb * 1024 * 1024:
                logger.warning(f"File too large, skipping: {file_path}")
                self._stats['files_skipped'] += 1
                return None
            
            # Get file info
            file_info = self._get_file_info(file_path)
            
            # Check if format is supported
            if not self._is_supported_format(file_info['extension']):
                self._stats['files_skipped'] += 1
                return None
            
            # Extract content
            content = await self._extract_content(file_path, file_info)
            
            if not content or len(content) < self.min_content_length:
                self._stats['files_skipped'] += 1
                return None
            
            # Build document
            document = {
                'content': content,
                'file_path': file_path,
                'file_name': file_info['name'],
                'file_type': file_info['extension'],
                'file_size': file_info['size'],
                'last_modified': file_info['modified'],
                'encoding': file_info.get('encoding', self.default_encoding)
            }
            
            # Add metadata if enabled
            if self.extract_metadata:
                metadata = await self._extract_metadata(file_path, file_info, content)
                document['metadata'] = metadata
            
            self._stats['files_processed'] += 1
            self._stats['total_size_bytes'] += file_info['size']
            
            return document
            
        except Exception as e:
            error_msg = f"Failed to load file {file_path}: {e}"
            logger.error(error_msg)
            self._stats['errors'].append(error_msg)
            self._stats['files_skipped'] += 1
            return None
    
    async def _find_supported_files(self) -> List[str]:
        """Find all supported files in the base path"""
        try:
            file_paths = []
            
            def scan_directory(path: str):
                for item in os.listdir(path):
                    if item.startswith('.') and not self.include_hidden:
                        continue
                    
                    if any(pattern in item for pattern in self.exclude_patterns):
                        continue
                    
                    item_path = os.path.join(path, item)
                    
                    if os.path.isfile(item_path):
                        file_ext = Path(item_path).suffix.lower()
                        if self._is_supported_format(file_ext):
                            file_paths.append(item_path)
                    
                    elif os.path.isdir(item_path) and self.recursive:
                        scan_directory(item_path)
            
            # Run directory scanning in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, scan_directory, self.base_path)
            
            return file_paths
            
        except Exception as e:
            logger.error(f"Failed to find files: {e}")
            return []
    
    async def _process_files_batch(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """Process multiple files concurrently"""
        try:
            # Process files in batches to avoid overwhelming the system
            batch_size = 10
            all_documents = []
            
            for i in range(0, len(file_paths), batch_size):
                batch_paths = file_paths[i:i + batch_size]
                
                # Create tasks for batch
                tasks = [self.load_file(path) for path in batch_paths]
                
                # Execute batch
                batch_results = await asyncio.gather(*tasks, return_exceptions=True)
                
                # Collect valid documents
                for result in batch_results:
                    if isinstance(result, dict):
                        all_documents.append(result)
                    elif isinstance(result, Exception):
                        logger.error(f"Batch processing error: {result}")
            
            return all_documents
            
        except Exception as e:
            logger.error(f"Batch processing failed: {e}")
            return []
    
    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information"""
        try:
            stat = os.stat(file_path)
            path_obj = Path(file_path)
            
            return {
                'name': path_obj.name,
                'extension': path_obj.suffix.lower(),
                'size': stat.st_size,
                'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'created': datetime.fromtimestamp(stat.st_ctime).isoformat()
            }
            
        except Exception as e:
            logger.error(f"Failed to get file info for {file_path}: {e}")
            return {
                'name': os.path.basename(file_path),
                'extension': Path(file_path).suffix.lower(),
                'size': 0,
                'modified': datetime.now().isoformat(),
                'created': datetime.now().isoformat()
            }
    
    def _is_supported_format(self, extension: str) -> bool:
        """Check if file format is supported"""
        return extension in self.supported_formats
    
    async def _extract_content(self, file_path: str, file_info: Dict[str, Any]) -> str:
        """Extract text content from file based on its type"""
        try:
            extension = file_info['extension']
            
            # Text files
            if extension in ['.txt', '.md', '.rst', '.log']:
                return await self._extract_text_file(file_path)
            
            # JSON files
            elif extension == '.json':
                return await self._extract_json_file(file_path)
            
            # CSV files
            elif extension == '.csv':
                return await self._extract_csv_file(file_path)
            
            # PDF files
            elif extension == '.pdf':
                return await self._extract_pdf_file(file_path)
            
            # Word documents
            elif extension in ['.docx', '.doc']:
                return await self._extract_word_file(file_path)
            
            # Excel files
            elif extension in ['.xlsx', '.xls']:
                return await self._extract_excel_file(file_path)
            
            # HTML files
            elif extension in ['.html', '.htm']:
                return await self._extract_html_file(file_path)
            
            # XML files
            elif extension == '.xml':
                return await self._extract_xml_file(file_path)
            
            # Python files
            elif extension == '.py':
                return await self._extract_code_file(file_path)
            
            # JavaScript files
            elif extension in ['.js', '.ts', '.jsx', '.tsx']:
                return await self._extract_code_file(file_path)
            
            # Other code files
            elif extension in ['.java', '.cpp', '.c', '.h', '.cs', '.php', '.rb', '.go', '.rs']:
                return await self._extract_code_file(file_path)
            
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Content extraction failed for {file_path}: {e}")
            return ""
    
    async def _extract_text_file(self, file_path: str) -> str:
        """Extract content from text files"""
        try:
            # Detect encoding if enabled
            encoding = self.default_encoding
            if self.detect_encoding:
                with open(file_path, 'rb') as f:
                    raw_data = f.read(10000)  # Read first 10KB for detection
                    detected = chardet.detect(raw_data)
                    if detected['encoding'] and detected['confidence'] > 0.7:
                        encoding = detected['encoding']
            
            # Read file content
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                content = f.read()
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return ""
    
    async def _extract_json_file(self, file_path: str) -> str:
        """Extract content from JSON files"""
        try:
            with open(file_path, 'r', encoding=self.default_encoding) as f:
                data = json.load(f)
            
            # Convert JSON to readable text
            return json.dumps(data, indent=2, ensure_ascii=False)
            
        except Exception as e:
            logger.error(f"Failed to extract JSON from {file_path}: {e}")
            return ""
    
    async def _extract_csv_file(self, file_path: str) -> str:
        """Extract content from CSV files"""
        try:
            # Read CSV
            df = pd.read_csv(file_path, encoding=self.default_encoding)
            
            # Convert to text representation
            content_parts = []
            
            # Add column headers
            content_parts.append("Columns: " + ", ".join(df.columns.tolist()))
            
            # Add sample data (first few rows)
            content_parts.append("\nSample data:")
            content_parts.append(df.head(10).to_string(index=False))
            
            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                content_parts.append("\nNumeric summary:")
                content_parts.append(df[numeric_cols].describe().to_string())
            
            return "\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"Failed to extract CSV from {file_path}: {e}")
            return ""
    
    async def _extract_pdf_file(self, file_path: str) -> str:
        """Extract content from PDF files"""
        try:
            content_parts = []
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            if self.preserve_structure:
                                content_parts.append(f"--- Page {page_num + 1} ---")
                            content_parts.append(text.strip())
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num + 1} from {file_path}: {e}")
            
            return "\n\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"Failed to extract PDF from {file_path}: {e}")
            return ""
    
    async def _extract_word_file(self, file_path: str) -> str:
        """Extract content from Word documents"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            # Extract table content
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(" | ".join(row_text))
                
                if table_text:
                    content_parts.append("\n".join(table_text))
            
            return "\n\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"Failed to extract Word document from {file_path}: {e}")
            return ""
    
    async def _extract_excel_file(self, file_path: str) -> str:
        """Extract content from Excel files"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content_parts = []
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if self.preserve_structure:
                        content_parts.append(f"--- Sheet: {sheet_name} ---")
                    
                    # Add sheet content
                    content_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
                    content_parts.append(df.head(10).to_string(index=False))
                    
                except Exception as e:
                    logger.warning(f"Failed to read sheet {sheet_name}: {e}")
            
            return "\n\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"Failed to extract Excel from {file_path}: {e}")
            return ""
    
    async def _extract_html_file(self, file_path: str) -> str:
        """Extract content from HTML files"""
        try:
            with open(file_path, 'r', encoding=self.default_encoding, errors='replace') as f:
                html_content = f.read()
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            logger.error(f"Failed to extract HTML from {file_path}: {e}")
            return ""
    
    async def _extract_xml_file(self, file_path: str) -> str:
        """Extract content from XML files"""
        try:
            with open(file_path, 'r', encoding=self.default_encoding, errors='replace') as f:
                xml_content = f.read()
            
            # Parse XML and extract text
            soup = BeautifulSoup(xml_content, 'xml')
            text = soup.get_text()
            
            # Clean up whitespace
            return ' '.join(text.split())
            
        except Exception as e:
            logger.error(f"Failed to extract XML from {file_path}: {e}")
            return ""
    
    async def _extract_code_file(self, file_path: str) -> str:
        """Extract content from code files"""
        try:
            with open(file_path, 'r', encoding=self.default_encoding, errors='replace') as f:
                content = f.read()
            
            # For code files, preserve the original formatting
            return content.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract code from {file_path}: {e}")
            return ""
    
    async def _extract_metadata(
        self, 
        file_path: str, 
        file_info: Dict[str, Any], 
        content: str
    ) -> Dict[str, Any]:
        """Extract metadata from document"""
        try:
            metadata = {
                'file_path': file_path,
                'file_name': file_info['name'],
                'file_type': file_info['extension'],
                'file_size': file_info['size'],
                'last_modified': file_info['modified'],
                'content_length': len(content),
                'word_count': len(content.split()) if content else 0,
                'line_count': content.count('\n') + 1 if content else 0
            }
            
            # Add format-specific metadata
            extension = file_info['extension']
            
            if extension == '.pdf':
                # PDF-specific metadata could be added here
                pass
            elif extension in ['.docx', '.doc']:
                # Word document metadata could be added here
                pass
            elif extension in ['.csv', '.xlsx', '.xls']:
                # Spreadsheet metadata could be added here
                pass
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract metadata from {file_path}: {e}")
            return {}
    
    def _reset_stats(self):
        """Reset loading statistics"""
        self._stats = {
            'files_processed': 0,
            'files_skipped': 0,
            'total_size_bytes': 0,
            'processing_time': 0,
            'errors': []
        }
    
    def get_stats(self) -> Dict[str, Any]:
        """Get loading statistics"""
        return self._stats.copy()
    
    async def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return self.supported_formats.copy()
    
    async def test_file_support(self, file_path: str) -> Dict[str, Any]:
        """Test if a file can be processed"""
        try:
            if not os.path.exists(file_path):
                return {"supported": False, "reason": "File does not exist"}
            
            file_info = self._get_file_info(file_path)
            
            # Check file size
            if file_info['size'] > self.max_file_size_mb * 1024 * 1024:
                return {"supported": False, "reason": "File too large"}
            
            # Check format
            if not self._is_supported_format(file_info['extension']):
                return {"supported": False, "reason": "Unsupported format"}
            
            # Try to extract content
            content = await self._extract_content(file_path, file_info)
            
            if not content or len(content) < self.min_content_length:
                return {"supported": False, "reason": "No extractable content"}
            
            return {
                "supported": True,
                "file_info": file_info,
                "content_preview": content[:200] + "..." if len(content) > 200 else content
            }
            
        except Exception as e:
            return {"supported": False, "reason": f"Error: {str(e)}"}


# Convenience functions
async def load_documents(base_path: str = None, **kwargs) -> List[Dict[str, Any]]:
    """Convenience function to load documents"""
    loader = DocumentLoader(base_path, **kwargs)
    return await loader.load()


async def load_single_document(file_path: str, **kwargs) -> Optional[Dict[str, Any]]:
    """Convenience function to load a single document"""
    loader = DocumentLoader(**kwargs)
    return await loader.load_file(file_path)